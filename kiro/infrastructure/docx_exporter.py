"""Exporta Article/CustomerFAQ para .docx (Microsoft Word / Google Docs).

Após redesign da issue #15: SEM sections "Problema/Causa raiz/Solução"; SEM
nota interna com OPE-XXX. Estrutura espelha o padrão SUP — título, scope
note em itálico, sections (ou entries) como H1, e nada mais.

Use case: enquanto o time não tem permissão no Confluence, os drafts saem
como .docx pra serem subidos no Google Drive e revisados pela equipe via
comentários.
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from kiro.domain.models import ArticleDraft, Cluster, CustomerFAQ
from kiro.utils.branding import SIGNATURE


def article_to_docx(article: ArticleDraft, cluster: Cluster, output_path: Path) -> Path:
    """Renderiza Article como .docx (padrão SUP — issue #15).

    Estrutura:
    - Título
    - scope_note em itálico cinza (subtítulo)
    - Cada section como H1 com body markdown como parágrafos
    - Tags
    - Rodapé com slogan
    """
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # Título
    doc.add_heading(article.title, level=0)

    # Scope note (substituiu o subtítulo de "rascunho de N tickets")
    scope = doc.add_paragraph()
    scope_run = scope.add_run(article.scope_note)
    scope_run.italic = True
    scope_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Sections
    for section in article.sections:
        doc.add_heading(section.heading, level=1)
        for paragraph in _split_markdown_paragraphs(section.body):
            doc.add_paragraph(paragraph)

    # Tags
    if article.tags:
        tags_para = doc.add_paragraph()
        tags_para.add_run("Tags: ").bold = True
        tags_run = tags_para.add_run(", ".join(article.tags))
        tags_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Rodapé com marca
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(SIGNATURE)
    footer_run.italic = True
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def customer_faq_to_docx(faq: CustomerFAQ, cluster: Cluster, output_path: Path) -> Path:
    """Renderiza FAQ como .docx (padrão SUP — issue #15).

    Estrutura:
    - Título
    - scope_note em itálico cinza
    - Cada entry como H1 (pergunta) + parágrafos (resposta) + when_to_contact
    - Tags
    - Rodapé
    """
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    doc.add_heading(faq.title, level=0)

    scope = doc.add_paragraph()
    scope_run = scope.add_run(faq.scope_note)
    scope_run.italic = True
    scope_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    for entry in faq.entries:
        q_heading = doc.add_heading(entry.question, level=1)
        for run in q_heading.runs:
            run.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)

        for paragraph in _split_markdown_paragraphs(entry.answer):
            doc.add_paragraph(paragraph)

        if entry.when_to_contact:
            wtc_para = doc.add_paragraph()
            wtc_para.add_run("Quando contatar o suporte: ").bold = True
            wtc_para.add_run(entry.when_to_contact)

    if faq.tags:
        tags_para = doc.add_paragraph()
        tags_para.add_run("Tags: ").bold = True
        tags_run = tags_para.add_run(", ".join(faq.tags))
        tags_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(SIGNATURE)
    footer_run.italic = True
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def _split_markdown_paragraphs(text: str) -> list[str]:
    """Quebra markdown em parágrafos preservando bullets simples.

    Não interpreta toda a sintaxe — só separa por linha em branco e mantém
    bullets `- ` como prefixo. Suficiente pro nível de fidelidade do .docx
    (revisão pela chefe, não impressão final).
    """
    paragraphs: list[str] = []
    current: list[str] = []
    for line in text.split("\n"):
        stripped = line.rstrip()
        if not stripped:
            if current:
                paragraphs.append("\n".join(current))
                current = []
            continue
        current.append(stripped)
    if current:
        paragraphs.append("\n".join(current))
    return paragraphs or [text]
